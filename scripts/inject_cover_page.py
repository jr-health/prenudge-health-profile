"""
Replaces the auto-generated Title/Date paragraphs that pandoc puts at the top of a
health-profile .docx with the real corporate cover page from the reference template
(render/templates/PräNUDGE Berichtsvorlage.docx), substituting in the real title,
version and generation date - and fills in the matching version/date placeholders in
both footers (the title-page-only footer and the default footer used on every other
page).

Runs after the asciidoctor -> docbook -> pandoc conversion (see release.yml, Phase 4
in doc/github-actions-plan.md). The reference .docx's cover page is a single paragraph
holding a logo picture plus two content controls ("Titel", "Autor"); each control has
an mc:AlternateContent Choice/Fallback pair (standard OOXML for any Word text box, not
a duplicate to remove) - both branches get the substituted text so the result is
correct however it's viewed. The title-page footer's date was a live DATE field
(always shows "today"); it gets replaced with the fixed `--generated` value so it
doesn't silently change every time someone opens the file.

Usage:
    python scripts/inject_cover_page.py health-profile-v1.2.0-2026-08-26.de.docx \
        --reference "render/templates/PräNUDGE Berichtsvorlage.docx" \
        --title "PreNUDGE Gesundheitsprofil — Katalog" \
        --version 1.2.0 --generated 2026-08-26 --lang de

Requires:
    pip install python-docx
"""

import argparse
import copy
import io
from pathlib import Path

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

VERSION_LINE = {
    "de": "Version {version} · Generiert: {generated}",
    "en": "Version {version} · Generated: {generated}",
}

CONSORTIUM_LABEL = "The PreNUDGE Consortium"
CONSORTIUM_URL = "https://prenudge.at/#konsortium"


def find_cover_paragraph(reference_doc):
    """The reference doc's cover page is body[0], a single w:sdt wrapping one
    paragraph that carries the logo picture and the Titel/Autor content controls."""
    cover_sdt = reference_doc.element.body[0]
    sdt_content = cover_sdt.find(qn("w:sdtContent"))
    for p in sdt_content.findall(qn("w:p")):
        if p.find(".//" + qn("w:drawing")) is not None:
            return p
    raise RuntimeError("No paragraph with a drawing found in the reference cover page (body[0]).")


def alias_of(sdt):
    sdt_pr = sdt.find(qn("w:sdtPr"))
    alias = sdt_pr.find(qn("w:alias")) if sdt_pr is not None else None
    return alias.get(qn("w:val")) if alias is not None else None


def set_sdt_text(sdt, text):
    """Collapse every <w:t> run inside the content control down to a single run."""
    runs = sdt.findall(".//" + qn("w:t"))
    if not runs:
        return
    runs[0].text = text
    for extra in runs[1:]:
        extra.text = ""


def make_text_run(text, rPr):
    r = OxmlElement("w:r")
    if rPr is not None:
        r.append(copy.deepcopy(rPr))
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def set_sdt_two_lines(sdt, line1_text, line2_text, line2_rid):
    """Replace the content control's paragraph with two lines: plain text, a line
    break, then a hyperlinked run - reusing the run formatting (color/bold/caps/size)
    already defined on the control's first existing run."""
    runs = sdt.findall(".//" + qn("w:r"))
    if not runs:
        return
    template_rPr = runs[0].find(qn("w:rPr"))
    para = runs[0].getparent()
    for r in list(para.findall(qn("w:r"))):
        para.remove(r)

    para.append(make_text_run(line1_text, template_rPr))

    br_run = OxmlElement("w:r")
    if template_rPr is not None:
        br_run.append(copy.deepcopy(template_rPr))
    br_run.append(OxmlElement("w:br"))
    para.append(br_run)

    link_rPr = copy.deepcopy(template_rPr) if template_rPr is not None else OxmlElement("w:rPr")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    link_rPr.append(underline)
    link_run = make_text_run(line2_text, link_rPr)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), line2_rid)
    hyperlink.append(link_run)
    para.append(hyperlink)


def unwrap_sdt(sdt):
    """Replace the content control with its own content, removing the control
    wrapper - including its <w:dataBinding> (bound to the reference doc's own
    customXml part, which doesn't exist in the pandoc-generated package; leaving
    it in place makes Word treat the file as unreadable on open)."""
    parent = sdt.getparent()
    idx = list(parent).index(sdt)
    sdt_content = sdt.find(qn("w:sdtContent"))
    for offset, child in enumerate(list(sdt_content)):
        parent.insert(idx + offset, child)
    parent.remove(sdt)


def replace_field_with_text(paragraph_element, text):
    """Replace a begin/instrText/separate/cached-result/end field run sequence (e.g.
    a live DATE field) with one plain static run - so the value stays fixed forever
    instead of re-evaluating to "today" every time the file is opened."""
    runs = paragraph_element.findall(qn("w:r"))
    begin_idx = end_idx = None
    for i, r in enumerate(runs):
        fld = r.find(qn("w:fldChar"))
        if fld is None:
            continue
        fld_type = fld.get(qn("w:fldCharType"))
        if fld_type == "begin" and begin_idx is None:
            begin_idx = i
        elif fld_type == "end":
            end_idx = i
    if begin_idx is None or end_idx is None:
        return
    field_runs = runs[begin_idx : end_idx + 1]
    template_rPr = None
    for r in field_runs:
        if r.find(qn("w:t")) is not None:
            template_rPr = r.find(qn("w:rPr"))
    field_runs[0].addprevious(make_text_run(text, template_rPr))
    for r in field_runs:
        paragraph_element.remove(r)


def replace_run_text(paragraph_element, old_text, new_text, drop_highlight=False):
    for r in paragraph_element.findall(qn("w:r")):
        t = r.find(qn("w:t"))
        if t is None or t.text != old_text:
            continue
        t.text = new_text
        if drop_highlight:
            rPr = r.find(qn("w:rPr"))
            highlight = rPr.find(qn("w:highlight")) if rPr is not None else None
            if highlight is not None:
                rPr.remove(highlight)
        return True
    return False


def update_footers(out_doc, version, generated, lang):
    """The reference doc's footer already carries a "Version [x.y] <date>" line on the
    title page (footer2 / first_page_footer) - but the date there is a live DATE field
    (always "today"), and the default footer (footer1, all other pages) has no version
    at all and just a literal "TT-MMM-JJJJ" placeholder someone left unfilled."""
    version_line = VERSION_LINE[lang].format(version=version, generated=generated)
    sec = out_doc.sections[0]

    for p in sec.footer.paragraphs:
        replace_run_text(p._p, "TT-MMM-JJJJ", version_line, drop_highlight=True)

    for p in sec.first_page_footer.paragraphs:
        replace_run_text(p._p, "Version [", f"Version {version}")
        replace_run_text(p._p, "x.y", "")
        replace_run_text(p._p, "]", "")
        replace_field_with_text(p._p, generated)


def strip_pandoc_title_paragraphs(out_doc):
    """Remove pandoc's auto-generated Title/Subtitle/Date paragraphs docbook's
    <info><title>/<date> produce at the very start of the body - they're being
    replaced by the real cover page."""
    body = out_doc.element.body
    auto_styles = {"Titel", "Title", "Untertitel", "Subtitle", "Date"}
    for p in list(body.findall(qn("w:p"))):
        pPr = p.find(qn("w:pPr"))
        style = pPr.find(qn("w:pStyle")) if pPr is not None else None
        style_val = style.get(qn("w:val")) if style is not None else None
        if style_val in auto_styles:
            body.remove(p)
        else:
            break  # stop at the first paragraph that isn't an auto title/date paragraph


def inject_cover_page(docx_path, reference_path, title, version, generated, lang):
    out_doc = docx.Document(str(docx_path))
    ref_doc = docx.Document(str(reference_path))

    cover_para = copy.deepcopy(find_cover_paragraph(ref_doc))
    version_line = VERSION_LINE[lang].format(version=version, generated=generated)
    consortium_rid = out_doc.part.relate_to(CONSORTIUM_URL, RT.HYPERLINK, is_external=True)
    for sdt in cover_para.findall(".//" + qn("w:sdt")):
        alias = alias_of(sdt)
        if alias == "Titel":
            set_sdt_text(sdt, title)
            unwrap_sdt(sdt)
        elif alias == "Autor":
            set_sdt_two_lines(sdt, version_line, CONSORTIUM_LABEL, consortium_rid)
            unwrap_sdt(sdt)

    for blip in cover_para.findall(".//" + qn("a:blip")):
        old_rid = blip.get(qn("r:embed"))
        if not old_rid:
            continue
        image_part = ref_doc.part.related_parts[old_rid]
        new_rid, _ = out_doc.part.get_or_add_image(io.BytesIO(image_part.blob))
        blip.set(qn("r:embed"), new_rid)

    strip_pandoc_title_paragraphs(out_doc)
    out_doc.element.body.insert(0, cover_para)
    update_footers(out_doc, version, generated, lang)

    out_doc.save(str(docx_path))


def main():
    parser = argparse.ArgumentParser(description="Inject the corporate cover page into a rendered health profile .docx")
    parser.add_argument("docx_path", help="Path to the pandoc-generated .docx (edited in place)")
    parser.add_argument("--reference", required=True, help="Path to the reference .docx carrying the cover page")
    parser.add_argument("--title", required=True, help="Report title")
    parser.add_argument("--version", required=True, help="Version string, e.g. 1.2.0")
    parser.add_argument("--generated", required=True, help="Generation date, e.g. 2026-08-26")
    parser.add_argument("--lang", required=True, choices=["de", "en"])
    args = parser.parse_args()

    inject_cover_page(Path(args.docx_path), Path(args.reference), args.title, args.version, args.generated, args.lang)
    print(f"  Cover page injected: {args.docx_path}")


if __name__ == "__main__":
    main()
